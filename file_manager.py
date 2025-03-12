import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches

class FileManager:
    def __init__(self, username):
        self.username = username
        self.base_path = "C:/NetworkScanner/Users"
        self.user_path = os.path.join(self.base_path, username)
        self.scans_path = os.path.join(self.user_path, "scans")
        self.reports_path = os.path.join(self.user_path, "reports")
        self.setup_directories()

    def setup_directories(self):
        """Create necessary directories if they don't exist"""
        for path in [self.user_path, self.scans_path, self.reports_path]:
            if not os.path.exists(path):
                os.makedirs(path)

    def save_scan_excel(self, scan_data, scan_id):
        """Save scan data to Excel file with enhanced formatting"""
        filename = f"scan_{scan_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        filepath = os.path.join(self.scans_path, filename)
        
        # Convert scan data to DataFrame with proper structure
        if isinstance(scan_data, dict):
            scan_data = [scan_data]
        
        df = pd.DataFrame(scan_data)
        
        # Create Excel writer object with xlsxwriter engine
        writer = pd.ExcelWriter(filepath, engine='openpyxl')
        
        # Write scan data
        df.to_excel(writer, sheet_name='Scan Results', index=False)
        
        # Get the xlsxwriter workbook and worksheet objects
        workbook = writer.book
        worksheet = writer.sheets['Scan Results']
        
        # Save the workbook
        writer.close()
        
        return filepath

    def save_report(self, doc, report_name):
        """Save a Word document report"""
        filepath = os.path.join(self.reports_path, f"{report_name}.docx")
        doc.save(filepath)
        return filepath

    def save_scan_word(self, scan_data, scan_id):
        """Save scan data to Word file with enhanced formatting"""
        filename = f"scan_{scan_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.reports_path, filename)
        
        doc = Document()
        doc.add_heading(f'Network Scan Report - {scan_id}', 0)
        
        # Add scan information
        doc.add_heading('Scan Details', level=1)
        doc.add_paragraph(f"Scan Time: {scan_data['scan_timestamp']}")
        doc.add_paragraph(f"IP Address: {scan_data['ip_address']}")
        doc.add_paragraph(f"Scan Type: {scan_data['scan_type']}")
        doc.add_paragraph(f"Status: {scan_data['status']}")
        
        # Add ports information if available
        if scan_data.get('ports'):
            doc.add_heading('Ports', level=1)
            doc.add_paragraph(f"Open Ports: {scan_data['ports']}")
        
        # Add services information if available
        if scan_data.get('services'):
            doc.add_heading('Services', level=1)
            doc.add_paragraph(f"Detected Services: {scan_data['services']}")
        
        # Add vulnerability information if available
        if any(key in scan_data for key in ['critical_vulns', 'high_vulns', 'medium_vulns', 'low_vulns']):
            doc.add_heading('Vulnerabilities', level=1)
            if scan_data.get('critical_vulns', 0) > 0:
                doc.add_paragraph(f"Critical Vulnerabilities: {scan_data['critical_vulns']}")
            if scan_data.get('high_vulns', 0) > 0:
                doc.add_paragraph(f"High Vulnerabilities: {scan_data['high_vulns']}")
            if scan_data.get('medium_vulns', 0) > 0:
                doc.add_paragraph(f"Medium Vulnerabilities: {scan_data['medium_vulns']}")
            if scan_data.get('low_vulns', 0) > 0:
                doc.add_paragraph(f"Low Vulnerabilities: {scan_data['low_vulns']}")
        
        doc.save(filepath)
        return filepath

    def generate_user_report(self, scan_history):
        """Generate comprehensive user report with enhanced formatting"""
        filename = f"user_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.reports_path, filename)
        
        doc = Document()
        doc.add_heading(f'Security Analysis Report - {self.username}', 0)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        summary.add_run(f"Total Scans Analyzed: {len(scan_history)}\n")
        
        # Vulnerability Summary
        if scan_history:
            total_critical = sum(scan.get('critical_vulns', 0) for scan in scan_history)
            total_high = sum(scan.get('high_vulns', 0) for scan in scan_history)
            total_medium = sum(scan.get('medium_vulns', 0) for scan in scan_history)
            total_low = sum(scan.get('low_vulns', 0) for scan in scan_history)
            
            doc.add_heading('Vulnerability Summary', level=1)
            vuln_summary = doc.add_paragraph()
            if total_critical > 0:
                vuln_summary.add_run(f"Critical Vulnerabilities: {total_critical}\n").bold = True
            if total_high > 0:
                vuln_summary.add_run(f"High Vulnerabilities: {total_high}\n").bold = True
            if total_medium > 0:
                vuln_summary.add_run(f"Medium Vulnerabilities: {total_medium}\n")
            if total_low > 0:
                vuln_summary.add_run(f"Low Vulnerabilities: {total_low}\n")
        
        # Recent Scans
        doc.add_heading('Recent Scans', level=1)
        for i, scan in enumerate(scan_history[-5:], 1):  # Last 5 scans
            doc.add_heading(f"Scan #{i}", level=2)
            scan_details = doc.add_paragraph()
            scan_details.add_run(f"Time: {scan['scan_timestamp']}\n")
            scan_details.add_run(f"IP: {scan['ip_address']}\n")
            scan_details.add_run(f"Type: {scan['scan_type']}\n")
            scan_details.add_run(f"Status: {scan['status']}\n")
            
            if scan.get('ports'):
                scan_details.add_run(f"Ports: {scan['ports']}\n")
            if scan.get('services'):
                scan_details.add_run(f"Services: {scan['services']}\n")
            if scan.get('os_info'):
                scan_details.add_run(f"OS Info: {scan['os_info']}\n")
            
            # Add vulnerability details
            if any(key in scan for key in ['critical_vulns', 'high_vulns', 'medium_vulns', 'low_vulns']):
                vuln_details = doc.add_paragraph()
                vuln_details.add_run("Vulnerabilities Found:\n").bold = True
                if scan.get('critical_vulns', 0) > 0:
                    vuln_details.add_run(f"- Critical: {scan['critical_vulns']}\n")
                if scan.get('high_vulns', 0) > 0:
                    vuln_details.add_run(f"- High: {scan['high_vulns']}\n")
                if scan.get('medium_vulns', 0) > 0:
                    vuln_details.add_run(f"- Medium: {scan['medium_vulns']}\n")
                if scan.get('low_vulns', 0) > 0:
                    vuln_details.add_run(f"- Low: {scan['low_vulns']}\n")
            
            doc.add_paragraph("---")
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        if total_critical > 0 or total_high > 0:
            doc.add_paragraph("⚠️ Immediate action required for critical and high vulnerabilities")
        doc.add_paragraph("• Regular security scans recommended")
        doc.add_paragraph("• Keep systems and software up to date")
        doc.add_paragraph("• Monitor network traffic for suspicious activities")
        
        doc.save(filepath)
        return filepath 