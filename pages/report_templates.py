import streamlit as st
from docx import Document
import datetime
from file_manager import FileManager
import pandas as pd

def init_style():
    """Initialize custom styling"""
    with open('static/style.css') as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

def create_network_scan_report(scan_results):
    doc = Document()
    doc.add_heading('Network Security Scan Report', 0)
    doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add scan summary
    doc.add_heading('Scan Summary', level=1)
    for host in scan_results:
        doc.add_heading(f'Host: {host["IP Address"]}', level=2)
        doc.add_paragraph(f'Status: {host["Status"]}')
        doc.add_paragraph(f'Hostname: {host["Hostname"]}')
        doc.add_paragraph(f'Operating System: {host["OS"]}')
        
        if host.get('Ports'):
            doc.add_paragraph(f'Open Ports: {", ".join(map(str, host["Ports"]))}')
        
        if host.get('Services'):
            doc.add_heading('Services', level=3)
            for service in host['Services']:
                doc.add_paragraph(f'• {service}')
        
        if host.get('Vulnerabilities'):
            doc.add_heading('Potential Vulnerabilities', level=3)
            for vuln in host['Vulnerabilities']:
                doc.add_paragraph(f'• {vuln}')
        
        doc.add_paragraph('---')
    
    return doc

def create_network_monitor_report(packet_data):
    doc = Document()
    doc.add_heading('Network Monitor Report', 0)
    doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add traffic summary
    doc.add_heading('Traffic Summary', level=1)
    
    # Group packets by protocol
    protocols = {}
    alerts = []
    
    for packet in packet_data:
        protocol = packet.get('protocol', 'Unknown')
        protocols[protocol] = protocols.get(protocol, 0) + 1
        
        if packet.get('alerts'):
            alerts.extend(packet['alerts'])
    
    # Add protocol statistics
    doc.add_heading('Protocol Distribution', level=2)
    for protocol, count in protocols.items():
        doc.add_paragraph(f'• {protocol}: {count} packets')
    
    # Add security alerts
    if alerts:
        doc.add_heading('Security Alerts', level=2)
        for alert in alerts:
            doc.add_paragraph(f'• Type: {alert["type"]}')
            doc.add_paragraph(f'  Source: {alert["source"]}')
            doc.add_paragraph(f'  Severity: {alert["severity"]}')
            doc.add_paragraph(f'  Description: {alert["description"]}')
            doc.add_paragraph('---')
    
    return doc

def display_report_templates():
    init_style()
    st.title("Report Templates")
    
    template_type = st.selectbox(
        "Select Report Type",
        ["Network Scan Report", "Network Monitor Report", "Executive Summary", "Technical Report", "Compliance Report"]
    )
    
    # Add report customization options
    st.markdown("### Report Options")
    
    include_graphs = st.checkbox("Include Graphs and Charts", value=True)
    include_raw_data = st.checkbox("Include Raw Data", value=False)
    
    if template_type == "Network Scan Report":
        st.markdown("""
        <div class='report-preview'>
            <h3>Network Scan Report Template</h3>
            <p>Includes:</p>
            <ul>
                <li>Host discovery results</li>
                <li>Open ports and services</li>
                <li>Vulnerability assessment</li>
                <li>OS detection results</li>
                <li>Network topology</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    elif template_type == "Network Monitor Report":
        st.markdown("""
        <div class='report-preview'>
            <h3>Network Monitor Report Template</h3>
            <p>Includes:</p>
            <ul>
                <li>Traffic analysis</li>
                <li>Security alerts</li>
                <li>Protocol distribution</li>
                <li>Suspicious activities</li>
                <li>Network performance metrics</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    elif template_type == "Executive Summary":
        st.markdown("""
        <div class='report-preview'>
            <h3>Executive Summary Template</h3>
            <p>Includes:</p>
            <ul>
                <li>Overview of scan results</li>
                <li>Key findings and risks</li>
                <li>Recommendations</li>
                <li>Summary metrics</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    # Generate Report Button
    if st.button("Generate Report", type="primary"):
        try:
            file_manager = FileManager(st.session_state.username)
            
            if template_type == "Network Monitor Report":
                # Generate network monitor report
                doc = create_network_monitor_report(st.session_state.get('packet_data', []))
                report_path = file_manager.save_report(doc, f"network_monitor_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}")
                st.success(f"Network Monitor Report generated successfully! Saved to: {report_path}")
            
            elif template_type == "Network Scan Report":
                # Generate network scan report
                doc = create_network_scan_report(st.session_state.get('scan_results', []))
                report_path = file_manager.save_report(doc, f"network_scan_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}")
                st.success(f"Network Scan Report generated successfully! Saved to: {report_path}")
            
            else:
                st.info("This report type is coming soon!")
        
        except Exception as e:
            st.error(f"Error generating report: {str(e)}")

    # Add custom CSS
    st.markdown("""
    <style>
    .report-preview {
        background: rgba(1, 1, 43, 0.7);
        padding: 1.5rem;
        border-radius: 15px;
        border: 1px solid var(--primary-color);
        margin: 1rem 0;
    }
    
    .report-preview h3 {
        color: var(--secondary-color);
        margin-bottom: 1rem;
    }
    
    .report-preview ul {
        margin-left: 1.5rem;
        margin-bottom: 1rem;
    }
    
    .report-preview li {
        margin-bottom: 0.5rem;
        color: var(--text-light);
    }
    </style>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    display_report_templates() 